"""
docx_render.py — python-docx ATS 兼容副路径
================================================================
HTML/PDF 是给人看的主版；docx 是给 ATS 解析的纯文本备份。
两路径共享 content dict（build.py assemble_content 输出），互不耦合。

设计原则（不要破坏）：
  - 极简：无表格、无文本框、无图片、无嵌入字体、无 page break
  - 结构化：Heading 1（name）/ Heading 2（section）/ Normal / List Bullet
  - 字面准确：所有内容来自 content dict，不补字、不删字
  - 视觉不追求跟 PDF 一致（CLAUDE.md 同步检查清单不要求一致）

ATS 兼容要点：
  - 联系信息单段：name + 多行 contact 走单独段，不进表格
  - 章节用 Heading 2 让 ATS 识别"实验性"区段（Workday/Greenhouse 用 heading detect）
  - bullet 用 List Bullet style，让 ATS 识别"this is a bulleted list"
  - title / date / org 分行，不要塞 tab right-align（部分 ATS 把 tab 当列分隔）

依赖：
  python3 -m pip install python-docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("缺少 python-docx，请先 python3 -m pip install python-docx")


HERE = Path(__file__).resolve().parent


# ============================================================
# 极简样式定义（不依赖外部模板）
# ============================================================

def _setup_styles(doc):
    """配置基础样式：字体一致（系统 serif），section heading 加粗。"""
    styles = doc.styles
    # Normal: 默认 body
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10)
    # Heading 1（名字）
    h1 = styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    # Heading 2（section）
    h2 = styles["Heading 2"]
    h2.font.name = "Georgia"
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def _set_margins(doc, top_in=0.6, bottom_in=0.6, left_in=0.7, right_in=0.7):
    """简历单页紧凑 margin。"""
    for section in doc.sections:
        section.top_margin = Cm(top_in * 2.54)
        section.bottom_margin = Cm(bottom_in * 2.54)
        section.left_margin = Cm(left_in * 2.54)
        section.right_margin = Cm(right_in * 2.54)


# ============================================================
# 简历渲染
# ============================================================

def render_resume(content, out_docx_path):
    """content: build.py assemble_content() 返回的 dict
    out_docx_path: docx 落地路径
    返回 out_docx_path。
    """
    doc = Document()
    _setup_styles(doc)
    _set_margins(doc)

    # ---- name + contact 头部（单段，居中） ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content["contact"]["name"])
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content["contact"]["line"])
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

    # ---- SUMMARY ----
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(content["summary"])

    # ---- EXPERIENCE ----
    doc.add_heading("Experience", level=2)
    for exp in content["experience"]:
        # 职位 + 日期 一行（title 加粗，date 普通；用空格分隔不用 tab）
        p = doc.add_paragraph()
        run_title = p.add_run(exp["title"])
        run_title.bold = True
        run_title.font.size = Pt(10.5)
        # 日期跟在职位后，分隔符用 "  ·  " 避免 ATS 把 tab 当列分隔
        p.add_run("  ·  ")
        run_date = p.add_run(exp["date"])
        run_date.italic = True
        run_date.font.size = Pt(9.5)
        run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # org_line 单独一段
        p = doc.add_paragraph()
        run_org = p.add_run(exp["org_line"])
        run_org.font.size = Pt(9.5)
        run_org.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # bullets 用 List Bullet
        for b in exp["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    # ---- EDUCATION ----
    doc.add_heading("Education", level=2)
    for edu in content["education"]:
        p = doc.add_paragraph()
        run_deg = p.add_run(edu["degree"])
        run_deg.bold = True
        run_deg.font.size = Pt(10.5)
        p.add_run("  ·  ")
        run_date = p.add_run(edu["date"])
        run_date.italic = True
        run_date.font.size = Pt(9.5)
        run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p = doc.add_paragraph()
        run_org = p.add_run(edu["org_line"])
        run_org.font.size = Pt(9.5)
        run_org.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- SKILLS ----
    doc.add_heading("Skills", level=2)
    for sk in content["skills"]:
        p = doc.add_paragraph()
        run_label = p.add_run(sk["label"] + " ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        p.add_run(sk["body"])

    doc.save(str(out_docx_path))
    return out_docx_path


# ============================================================
# Cover letter 渲染
# ============================================================

def render_cl(content, cl_data, out_docx_path):
    """CL docx 渲染。
    content: 跟 resume 同一个 dict（取 contact header）
    cl_data: APP###.yaml 的 cover_letter 段
    """
    doc = Document()
    _setup_styles(doc)
    _set_margins(doc, top_in=0.8, bottom_in=0.8, left_in=0.85, right_in=0.85)

    # ---- name + contact header（跟简历一致） ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content["contact"]["name"])
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content["contact"]["line"])
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

    # ---- date ----
    doc.add_paragraph()  # 空行
    doc.add_paragraph(cl_data["date"])

    # ---- recipient_company ----
    p = doc.add_paragraph()
    run = p.add_run(cl_data["recipient_company"])
    run.bold = True

    # ---- Re ----
    p = doc.add_paragraph()
    run = p.add_run(f"Re: {cl_data['re_line']}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # 空行

    # ---- salutation ----
    doc.add_paragraph(cl_data["salutation"])

    # ---- body_paragraphs ----
    for para in cl_data.get("body_paragraphs", []):
        if para and para.strip():
            doc.add_paragraph(para)

    # ---- closing + signature ----
    doc.add_paragraph()  # 空行
    doc.add_paragraph(cl_data["closing"])
    doc.add_paragraph(cl_data["signature"])

    doc.save(str(out_docx_path))
    return out_docx_path


# ============================================================
# CLI 自测
# ============================================================

if __name__ == "__main__":
    import yaml as _yaml
    repo_root = HERE.parent
    master = _yaml.safe_load(open(repo_root / "master_resume.yaml", encoding="utf-8"))

    content = {
        "contact": master["contact"],
        "summary": master["summary"],
        "experience": [
            {"title": e["title"], "date": e["date"], "org_line": e["org_line"], "bullets": e["bullets"]}
            for e in master["experience"]
        ],
        "education": master["education"],
        "skills": [{"label": s["label"], "body": s["body"]} for s in master["skills"]],
    }

    out_dir = HERE / "applications" / "_spike"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / "docx_render_selftest.docx"

    print(f"  渲染 resume docx ...")
    render_resume(content, out_docx)
    print(f"  ✓ {out_docx} ({out_docx.stat().st_size/1024:.1f} KB)")
